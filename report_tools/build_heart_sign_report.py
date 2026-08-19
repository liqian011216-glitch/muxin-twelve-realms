from __future__ import annotations

from datetime import date
from itertools import product
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/yizhixiaojinli/牧心十二境")
OUTPUT = ROOT / "交付文档" / "牧心十二境_81种心签算法与内容报告.docx"

BLUE = "284C63"
DARK_BLUE = "17374A"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7F9"
RED = "A63A2B"
GOLD = "A67C2E"
MUTED = "66737B"
GRID = "B7C4CC"
BLACK = "1F2529"

# Keep the report visually consistent with the game. The matching OTF is also
# supplied to LibreOffice through SAL_FONTPATH during render verification.
ASCII_FONT = "Huiwen Mingchao"
CJK_FONT = "Huiwen Mingchao"
CJK_SERIF = "Huiwen Mingchao"


PATHS = {
    "explore": {
        "choice": "未经驯服（untrained）",
        "title": "逐野",
        "verse": "野径初开随迹去。",
        "heart": "你先看见未经驯服的力量，",
        "tie": "行动 → 平衡 → 观照",
    },
    "hold": {
        "choice": "受绳约束（restrained）",
        "title": "持衡",
        "verse": "手持一线过重关。",
        "heart": "你先看见约束中的秩序，",
        "tie": "平衡 → 观照 → 行动",
    },
    "release": {
        "choice": "自在无绳（free）",
        "title": "任流",
        "verse": "松绳仍照水云间。",
        "heart": "你先看见自在之中的觉照，",
        "tie": "观照 → 平衡 → 行动",
    },
}

DOMINANTS = {
    "action": {
        "label": "行动",
        "title": "行",
        "verse": "几番追逐识心奔。",
        "heart": "心性所显，多在行动之中，",
    },
    "balance": {
        "label": "平衡",
        "title": "定",
        "verse": "桥心一步一安然。",
        "heart": "心性所显，多在调衡之中，",
    },
    "awareness": {
        "label": "观照",
        "title": "照",
        "verse": "花入衣襟知细微。",
        "heart": "心性所显，多在留意之中，",
    },
}

PACES = {
    "swift": {
        "label": "疾",
        "verse": "风起即行不滞身。",
        "heart": "临路即行，你回应得快，",
        "range": "莲花完成时间 ≤ 65秒",
    },
    "measured": {
        "label": "和",
        "verse": "缓急相宜自有程。",
        "heart": "缓急相宜，你自有节奏，",
        "range": "66–95秒",
    },
    "patient": {
        "label": "徐",
        "verse": "迟行亦把沿途照。",
        "heart": "你把时间，留给了过程，",
        "range": "≥ 96秒",
    },
}

DEPTHS = {
    "opening": {
        "label": "初见",
        "verse": "未圆之处仍堪问。",
        "heart": "此刻不必急着成为答案。",
        "range": "综合分 < 0.42",
    },
    "growing": {
        "label": "渐明",
        "verse": "回首已知来处路。",
        "heart": "你已在往复之间认出来路。",
        "range": "0.42 ≤ 综合分 < 0.72",
    },
    "clear": {
        "label": "通透",
        "verse": "月到中天照本真。",
        "heart": "所行渐明，心月自然显现。",
        "range": "综合分 ≥ 0.72",
    },
}

CATCH_LINES = [
    ("0–9次", "追牛未多，你仍在辨路，"),
    ("10–17次", "牛影数近，你能屡次回身，"),
    ("18次及以上", "屡次相逢，见你追寻不息，"),
]

BRIDGE_LINES = [
    ("26秒及以上", "桥上多摇，你仍谨慎修正，"),
    ("17–25秒", "桥有偏移，你能重归平衡，"),
    ("16秒以内", "过桥从容，步调很快安定，"),
]

FLOWER_LINES = [
    ("0–19朵", "花未尽收，目光更向前路，"),
    ("20–34朵", "沿途有花，常得你的照看，"),
    ("35–48朵", "细花入眼，多被你所看见，"),
]


def set_run_font(run, *, size=11, bold=False, color=BLACK, serif=False, italic=False):
    run.font.name = ASCII_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ASCII_FONT)
    fonts.set(qn("w:hAnsi"), ASCII_FONT)
    fonts.set(qn("w:eastAsia"), CJK_SERIF if serif else CJK_FONT)
    fonts.set(qn("w:cs"), CJK_SERIF if serif else CJK_FONT)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True
        fmt.keep_together = True


def set_paragraph_shading(paragraph, fill=PALE_BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_text(doc, text="", *, size=11, bold=False, color=BLACK, before=0, after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT, serif=False, italic=False, keep=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after, keep=keep)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, serif=serif, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def style_table_text(table, header=True, body_size=9.2):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=0, line=1.15)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.2 if row_idx == 0 and header else body_size,
                        bold=row_idx == 0 and header,
                        color=DARK_BLUE if row_idx == 0 and header else BLACK,
                    )
            if row_idx == 0 and header:
                shade_cell(cell, LIGHT_BLUE)
    if header:
        set_repeat_table_header(table.rows[0])


def add_key_value_table(doc, rows, widths=(2100, 7260)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = rows[0][0]
    table.rows[0].cells[1].text = rows[0][1]
    for left, right in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    set_table_geometry(table, widths)
    for row in table.rows:
        shade_cell(row.cells[0], PALE_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=BLACK)
        for cell in row.cells:
            set_paragraph_spacing(cell.paragraphs[0], after=0, line=1.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_standard_table(doc, headers, rows, widths, body_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    set_table_geometry(table, widths)
    style_table_text(table, body_size=body_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    end = paragraph.add_run(" 页")
    set_run_font(end, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = ASCII_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_SERIF)
        style._element.rPr.rFonts.set(qn("w:cs"), CJK_SERIF)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=0, line=1.0)
    hr = hp.add_run("牧心十二境｜心签算法设计报告")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    doc.core_properties.title = "《牧心十二境》81种心签算法与内容报告"
    doc.core_properties.subject = "五项游戏数据、规则引擎、81种心签内容及判断依据"
    doc.core_properties.author = "《牧心十二境》项目组"
    doc.core_properties.keywords = "牧心十二境, 心签, 游戏算法, 互动叙事, 行为数据"


def build_cover(doc):
    add_text(doc, "数字文化互动叙事 · 算法设计说明", size=11, bold=True, color=GOLD,
             after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "《牧心十二境》", size=28, bold=True, color=DARK_BLUE,
             serif=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "81种心签算法与内容报告", size=22, bold=True, color=BLUE,
             serif=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "五项真实互动数据 × 确定性规则引擎 × 文学化行为映照", size=12,
             color=MUTED, after=54, align=WD_ALIGN_PARAGRAPH.CENTER)

    metadata = [
        ("报告版本", "V1.0（与当前游戏实现一致）"),
        ("生成日期", "2026年8月19日"),
        ("算法类型", "可解释、确定性、非随机的规则型生成算法"),
        ("核心签型", "81种（3 × 3 × 3 × 3）"),
        ("数据来源", "问心选择、抓牛次数、过桥时间、莲花数量、莲花完成时间"),
    ]
    add_key_value_table(doc, metadata, widths=(2100, 7260))
    add_text(doc, "用途：项目答辩、算法说明、内容审核与后续用户测试校准。", size=9.5,
             color=MUTED, italic=True, before=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_summary(doc):
    add_heading(doc, "执行摘要", 1)
    p = add_text(
        doc,
        "本项目的心签不是随机抽签，也不是心理诊断。系统读取玩家在本次完整旅程中产生的五项真实数据，先把数据映射为问心路径、主导行为、行进节奏和完成层次，再由固定文案模块组合为签名、四句签诗与七句心解。相同输入必定得到相同输出，整个过程可以逐项解释和复算。",
        after=10,
    )
    set_paragraph_shading(p)

    add_key_value_table(doc, [
        ("算法结论", "3种路径 × 3种主导行为 × 3种节奏 × 3种层次 = 81种核心签型。"),
        ("显示差异", "81种签型拥有不同的内部ID或签诗组合；签名本身为27种，层次通过第四句签诗和第七句心解体现。"),
        ("动态心解", "七句心解中的第2、3、4句继续根据抓牛、过桥和莲花数量的表现分档变化，因此最终显示文本不止81个固定段落。"),
        ("数据持久化", "当前版本把五项数据保存在玩家当前浏览器的 sessionStorage 中，并在本次会话内生成结果。"),
        ("适用边界", "结果只用于叙事反馈和自我观察，不构成性格测验、临床评估或心理学诊断。"),
    ])

    add_heading(doc, "1. 报告范围与依据来源", 1)
    add_text(doc, "本报告依据当前可运行程序、自动化测试和已确认的实测阈值编制，目标是让指导教师能够回答三个问题：系统读了什么、怎样判断、最终文字如何产生。")
    add_standard_table(doc,
        ["依据类型", "具体来源", "在报告中的作用"],
        [
            ("程序实现", "app/player-sign.ts", "定义五项输入、评分、分档、并列处理和全部文案模块。"),
            ("数据管线", "app/page.tsx", "负责接收游戏数据并保存至浏览器会话。"),
            ("自动化验证", "tests/player-sign-algorithm.test.mjs", "验证81种结构、实测基准、边界值和缺失数据回退。"),
            ("实测校准", "抓牛18次；独木桥无掉落约16秒", "确定行动高表现参考值和过桥高表现边界。"),
            ("内容来源", "项目已确认的牧牛图禅理与互动叙事设定", "形成文学化签诗和心解；不是外部心理量表。"),
        ],
        [1600, 3000, 4760],
        body_size=9,
    )
    add_text(doc, "来源声明：心签文案为本项目规则型叙事文案，不是从古籍签文库随机抽取；算法也没有调用大模型实时判断玩家。", size=10, bold=True, color=RED, before=6)


def add_data_and_algorithm(doc):
    add_heading(doc, "2. 五项原始数据", 1)
    add_standard_table(doc,
        ["序号", "变量", "获取位置", "数据意义", "合法范围/单位"],
        [
            ("1", "cowChoice", "问心三选一", "决定叙事路径与并列分数的优先顺序", "untrained / restrained / free"),
            ("2", "caught", "寻牛游戏", "60秒内人物实际碰到牛的次数", "0次起，整数"),
            ("3", "bridgeSeconds", "独木桥游戏", "真正到岸时的累计用时；掉落后计时不重置", "秒"),
            ("4", "flowers", "莲花跳跃", "实际收集的莲花数量", "0–48朵"),
            ("5", "lotusSeconds", "莲花跳跃", "抵达终点时的累计完成时间", "秒"),
        ],
        [650, 1550, 1600, 3700, 1860],
        body_size=8.7,
    )

    add_heading(doc, "3. 判断算法与计算标准", 1)
    add_heading(doc, "3.1 三项标准化行为分", 2)
    add_standard_table(doc,
        ["维度", "计算公式", "解释"],
        [
            ("行动分 A", "A = clamp(抓牛次数 ÷ 18, 0, 1)", "18次及以上封顶为1；保留低中高之间的连续差异。"),
            ("平衡分 B", "≤16秒：1；17–25秒：0.65；≥26秒：0.30", "16秒来自无掉落实测；掉落造成的重走时间继续累计。"),
            ("观照分 C", "C = clamp(莲花数量 ÷ 48, 0, 1)", "48朵为全收集；数量越高代表本次路径中注意到更多细节。"),
        ],
        [1550, 3400, 4410],
        body_size=9,
    )

    add_heading(doc, "3.2 主导行为", 2)
    add_text(doc, "主导行为取 A、B、C 三项中的最高分：A最高为“行动”，B最高为“平衡”，C最高为“观照”。分数相同时不随机，而由问心选择规定优先顺序。")
    add_standard_table(doc,
        ["问心选择", "路径名", "并列时优先顺序"],
        [(value["choice"], value["title"], value["tie"]) for value in PATHS.values()],
        [3000, 1600, 4760],
        body_size=9.2,
    )

    add_heading(doc, "3.3 行进节奏", 2)
    add_standard_table(doc,
        ["节奏", "莲花完成时间", "算法含义"],
        [
            ("疾（swift）", "≤65秒", "响应快速；不直接等同于更优秀。"),
            ("和（measured）", "66–95秒", "速度与观察相对均衡。"),
            ("徐（patient）", "≥96秒", "愿意把时间留给过程；不作为失败。"),
        ],
        [1800, 2200, 5360],
        body_size=9.2,
    )

    add_heading(doc, "3.4 完成层次", 2)
    add_text(doc, "综合分 O = (A + B + C) ÷ 3。它只描述本次游戏中的完成程度，不是人格高低。")
    add_standard_table(doc,
        ["层次", "综合分范围", "叙事定位"],
        [
            ("初见（opening）", "O < 0.42", "仍在辨认自身的行动方式。"),
            ("渐明（growing）", "0.42 ≤ O < 0.72", "已能在往复中认出来路。"),
            ("通透（clear）", "O ≥ 0.72", "本次旅程中三项表现整体较完整。"),
        ],
        [2100, 2400, 4860],
        body_size=9.2,
    )

    add_heading(doc, "3.5 81种组合", 2)
    add_text(doc, "核心签型ID的结构为：路径 × 主导行为 × 节奏 × 层次。")
    p = add_text(doc, "3 × 3 × 3 × 3 = 81", size=18, bold=True, color=RED,
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=8)
    set_paragraph_shading(p, "F8EFEA")
    add_text(doc, "程序ID示例：explore-action-swift-clear。相同五项输入始终产生相同ID和相同文本，不使用随机数。")


def add_text_library(doc):
    add_heading(doc, "4. 文本生成结构与完整句库", 1)
    add_heading(doc, "4.1 输出结构", 2)
    add_standard_table(doc,
        ["输出部分", "生成规则"],
        [
            ("签名", "路径名 + ‘·’ + 主导字 + 节奏字 + ‘签’。例如：逐野·行疾签。"),
            ("四句签诗", "第1句取路径；第2句取主导行为；第3句取节奏；第4句取层次。"),
            ("七句心解", "第1句路径；第2句抓牛分档；第3句过桥分档；第4句莲花数量分档；第5句节奏；第6句主导行为；第7句层次。"),
        ],
        [1900, 7460],
        body_size=9.2,
    )

    add_heading(doc, "4.2 路径、主导、节奏、层次固定句库", 2)
    fixed_rows = []
    for key, value in PATHS.items():
        fixed_rows.append((f"路径｜{value['title']}", value["verse"], value["heart"]))
    for key, value in DOMINANTS.items():
        fixed_rows.append((f"主导｜{value['label']}", value["verse"], value["heart"]))
    for key, value in PACES.items():
        fixed_rows.append((f"节奏｜{value['label']}", value["verse"], value["heart"]))
    for key, value in DEPTHS.items():
        fixed_rows.append((f"层次｜{value['label']}", value["verse"], value["heart"]))
    add_standard_table(doc, ["模块", "签诗句", "心解句"], fixed_rows, [1800, 3400, 4160], body_size=8.7)

    add_heading(doc, "4.3 三项表现动态句库", 2)
    dynamic_rows = []
    for bucket, line in CATCH_LINES:
        dynamic_rows.append(("抓牛次数", bucket, line))
    for bucket, line in BRIDGE_LINES:
        dynamic_rows.append(("过桥时间", bucket, line))
    for bucket, line in FLOWER_LINES:
        dynamic_rows.append(("莲花数量", bucket, line))
    add_standard_table(doc, ["对应心解", "分档", "实际显示文字"], dynamic_rows, [1800, 2200, 5360], body_size=9)

    add_heading(doc, "4.4 缺失数据回退签", 2)
    add_text(doc, "只要五项数据中任一项为空，程序不会猜测，而是显示固定回退签《归径签》。")
    add_standard_table(doc,
        ["部分", "内容"],
        [
            ("签名", "归径签"),
            ("签诗", "路远不妨回首。／心驰亦可重收。／几番失足非为失。／云散之时月本明。"),
            ("心解", "你并非从未动摇，／而是每次动摇之后，／仍愿意重新回来。／路途虽曾反复，／只要心知归处，／不必催促此刻，／明月终会自现。"),
        ],
        [1600, 7760],
        body_size=9.2,
    )


def add_examples_and_validation(doc):
    add_heading(doc, "5. 计算示例", 1)
    examples = [
        (
            "示例A｜快速且整体完成度高",
            "问心：未经驯服；抓牛18次；过桥16秒；莲花48朵；莲花60秒",
            "A=1.00，B=1.00，C=1.00；并列按‘未经驯服’优先行动；节奏=疾；综合分=1.00，层次=通透。",
            "explore-action-swift-clear｜逐野·行疾签",
        ),
        (
            "示例B｜平衡主导、节奏适中",
            "问心：受绳约束；抓牛10次；过桥20秒；莲花24朵；莲花80秒",
            "A=0.56，B=0.65，C=0.50；主导=平衡；节奏=和；综合分≈0.57，层次=渐明。",
            "hold-balance-measured-growing｜持衡·定和签",
        ),
        (
            "示例C｜观照主导、徐行",
            "问心：自在无绳；抓牛3次；过桥30秒；莲花40朵；莲花110秒",
            "A=0.17，B=0.30，C=0.83；主导=观照；节奏=徐；综合分≈0.43，层次=渐明。",
            "release-awareness-patient-growing｜任流·照徐签",
        ),
        (
            "示例D｜数据不足",
            "任一数据未收到，例如抓牛次数为空",
            "不进行补值、不使用随机数、不推断缺失项。",
            "fallback-returning-path｜归径签",
        ),
    ]
    for title, inputs, calc, result in examples:
        h = add_heading(doc, title, 3)
        h.paragraph_format.keep_with_next = True
        p1 = add_text(doc, f"输入：{inputs}", size=9.8, keep=True)
        p2 = add_text(doc, f"计算：{calc}", size=9.8, keep=True)
        p3 = add_text(doc, f"结果：{result}", size=9.8, bold=True, color=RED, after=10)
        for p in (p1, p2):
            p.paragraph_format.keep_with_next = True

    add_heading(doc, "6. 可解释性、验证与限制", 1)
    add_key_value_table(doc, [
        ("确定性", "同一组五项数据重复计算，结果完全一致。"),
        ("全组合可达", "程序遍历验证表明81个路径—主导—节奏—层次ID均可由合法输入产生。"),
        ("边界明确", "抓牛10次进入中段、18次进入高段；过桥16秒进入高段、26秒进入低段；莲花20朵和35朵分别进入中段与高段。"),
        ("非随机", "算法不调用随机数，也不依赖联网服务或大模型实时生成。"),
        ("非诊断", "‘行动、平衡、观照、疾、和、徐’是叙事维度，不是经心理测量学验证的人格标签。"),
        ("待校准项", "莲花时间65/95秒、层次0.42/0.72、莲花20/35朵属于当前设计阈值，后续应根据更多真实玩家分布复核。"),
        ("隐私", "当前数据只保存在浏览器会话；若未来上传服务器，需要增加明示同意、数据最小化和删除机制。"),
    ])

    add_heading(doc, "7. 81种核心签型内容总览", 1)
    add_text(doc, "以下逐条列出当前程序中的81个核心签型。每条包含完整签名、四句固定签诗及七句心解的固定骨架。心解第2—4句必须按第4.3节的三张分档句库填入；因此同一核心签型也能真实反映玩家的抓牛、过桥和莲花表现。", after=12)


def add_sign_catalog(doc):
    combinations = list(product(PATHS.items(), DOMINANTS.items(), PACES.items(), DEPTHS.items()))
    for idx, ((path_key, path), (dom_key, dominant), (pace_key, pace), (depth_key, depth)) in enumerate(combinations, start=1):
        sign_id = f"{path_key}-{dom_key}-{pace_key}-{depth_key}"
        title = f"{path['title']}·{dominant['title']}{pace['label']}签"

        heading = add_heading(doc, f"{idx:02d}｜{title}", 3)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True

        basis = add_text(
            doc,
            f"ID：{sign_id}　｜　路径：{path['title']}　主导：{dominant['label']}　节奏：{pace['label']}　层次：{depth['label']}",
            size=8.6,
            color=MUTED,
            after=3,
            keep=True,
        )
        set_paragraph_shading(basis, PALE_BLUE)

        verse = add_text(
            doc,
            "签诗：" + "　".join((path["verse"], dominant["verse"], pace["verse"], depth["verse"])),
            size=9.2,
            serif=True,
            after=3,
            keep=True,
        )

        heart = add_text(
            doc,
            "心解骨架：①" + path["heart"]
            + "　②〔抓牛分档句〕　③〔过桥分档句〕　④〔莲花分档句〕　⑤"
            + pace["heart"] + "　⑥" + dominant["heart"] + "　⑦" + depth["heart"],
            size=8.8,
            after=8,
        )
        heart.paragraph_format.keep_together = True

        if idx % 3 == 0 and idx != len(combinations):
            heart.add_run().add_break(WD_BREAK.PAGE)


def add_appendix(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "附录A｜答辩时可直接使用的算法说明", 1)
    quote = add_text(
        doc,
        "“本项目的心签并非随机抽取，而是由玩家在问心、寻牛、独木桥和莲花跳跃中的五项真实数据共同决定。系统先把抓牛次数、过桥时间和莲花数量标准化为行动、平衡、观照三项得分，再结合问心选择确定路径和并列优先级；莲花完成时间用于判断节奏，三项得分的平均值用于判断完成层次。四层各三种，共形成81种核心签型。签文只对本次游戏行为做文学化映照，不作为心理诊断。”",
        size=11,
        color=DARK_BLUE,
        serif=True,
        before=8,
        after=14,
    )
    set_paragraph_shading(quote, LIGHT_BLUE)

    add_heading(doc, "附录B｜版本维护建议", 1)
    add_key_value_table(doc, [
        ("短期", "完成更多玩家测试，记录抓牛、过桥和莲花时间的真实分布。"),
        ("中期", "按分位数重新校准65/95秒、0.42/0.72等暂定阈值，并保留版本号。"),
        ("内容", "如果希望81种结果在标题层面也完全不同，可把‘初见／渐明／通透’加入签名；当前版本标题为27种、完整核心签型为81种。"),
        ("审计", "每次修改阈值或句库后，重新运行81组合覆盖测试、边界测试和缺失数据测试。"),
        ("隐私", "如增加排行榜或服务器存档，先明确告知采集项目、用途、保留期限和删除方式。"),
    ])

    add_heading(doc, "附录C｜当前实现文件", 1)
    add_standard_table(doc,
        ["文件", "作用"],
        [
            ("web/.worktrees/muxin-cover-step/app/player-sign.ts", "心签算法、阈值、文案模块与回退签。"),
            ("web/.worktrees/muxin-cover-step/app/page.tsx", "五项数据接收、会话存储与结局页调用。"),
            ("web/.worktrees/muxin-cover-step/tests/player-sign-algorithm.test.mjs", "算法单元测试与边界验证。"),
            ("web/.worktrees/muxin-cover-step/docs/心签算法依据.md", "项目内部简版算法说明。"),
        ],
        [5600, 3760],
        body_size=9,
    )


def build_document():
    doc = Document()
    configure_document(doc)
    build_cover(doc)
    add_summary(doc)
    add_data_and_algorithm(doc)
    add_text_library(doc)
    add_examples_and_validation(doc)
    add_sign_catalog(doc)
    add_appendix(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
