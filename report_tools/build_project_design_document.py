from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

import build_heart_sign_report as report_base

from build_heart_sign_report import (
    BLUE,
    DARK_BLUE,
    GOLD,
    MUTED,
    PALE_BLUE,
    RED,
    add_heading,
    add_key_value_table,
    add_page_number,
    add_standard_table,
    add_text,
    configure_document,
    set_paragraph_shading,
    set_paragraph_spacing,
    set_run_font,
)

# Use a system CJK font for the report so Word and the headless QA renderer both
# display Chinese reliably. The game itself continues to use Huiwen Mingchao.
report_base.ASCII_FONT = "Songti SC"
report_base.CJK_FONT = "Songti SC"
report_base.CJK_SERIF = "Songti SC"


ROOT = Path("/Users/yizhixiaojinli/牧心十二境")
WEB = ROOT / "web/.worktrees/muxin-cover-step"
OUTPUT = ROOT / "交付文档/牧心十二境_项目设计与技术实现说明书.docx"


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.15) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption_paragraph, after=10, line=1.0)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=8.8, color=MUTED)


def add_callout(doc: Document, text: str) -> None:
    paragraph = add_text(doc, text, size=10.5, color=DARK_BLUE, before=4, after=10)
    set_paragraph_shading(paragraph, PALE_BLUE)


def configure(doc: Document) -> None:
    configure_document(doc)
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].clear()
    header_run = header.paragraphs[0].add_run("牧心十二境｜项目设计与技术实现说明书")
    set_run_font(header_run, size=9, color=MUTED)
    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])
    doc.core_properties.title = "《牧心十二境》项目设计与技术实现说明书"
    doc.core_properties.subject = "用户需求、设计概念、信息架构、交互流程、视觉设计与技术实现"
    doc.core_properties.author = "《牧心十二境》项目组"
    doc.core_properties.keywords = "互动叙事, H5游戏, 牧牛图, 用户体验, Godot, React, 心签算法"


def build_cover(doc: Document) -> None:
    add_text(doc, "数字文化互动叙事项目", size=11, bold=True, color=GOLD,
             after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "《牧心十二境》", size=28, bold=True, color=DARK_BLUE,
             serif=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "项目设计与技术实现说明书", size=21, bold=True, color=BLUE,
             serif=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "从石刻文化认知到行为数据生成心签的手机横屏互动体验", size=11.5,
             color=MUTED, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure(doc, WEB / "public/frames/00-cover-muxin.png", "项目封面与整体视觉基调", 5.75)
    add_key_value_table(doc, [
        ("文档版本", "V1.0（依据2026年8月19日当前可运行版本编制）"),
        ("项目形态", "手机横屏H5互动叙事游戏，可通过公开网址访问"),
        ("核心内容", "牧牛十二境文化介绍、四阶段互动旅程、81种行为心签"),
        ("技术构成", "React / Next兼容前端 + Godot Web小游戏 + 浏览器本地数据管线"),
        ("适用场景", "毕业设计说明、课堂汇报、项目答辩、开发交接与后续测试"),
    ])
    page_break(doc)


def add_executive_summary(doc: Document) -> None:
    add_heading(doc, "执行摘要", 1)
    add_callout(doc, "《牧心十二境》以传统牧牛图为文化母题，将“观心—驯心—相忘—归真”的精神路径转化为可操作、可感知、可记录的手机横屏互动体验。玩家不是被动阅读十二幅图，而是通过拓印、追牛、长按牵绳、控制张力、过独木桥、收集莲花等行为完成一段由外在动作通向内在观照的旅程。")
    add_text(doc, "项目采用双层结构：前半部分负责建立文化背景和视觉语境，后半部分以四阶段游戏化流程承载体验。最终系统读取问心选择、抓牛次数、过桥时间、莲花数量和莲花完成时间五项真实数据，通过确定性规则生成81种心签之一，使结果既具有文学性，又能解释其判断来源。")
    add_key_value_table(doc, [
        ("核心命题", "牛既是画面中的对象，也是玩家妄心、习性与觉照状态的象征。"),
        ("体验目标", "让玩家在较低学习成本下理解牧牛图，并通过身体操作形成个人化记忆。"),
        ("设计原则", "不设置失败结局；失误转化为时间与过程数据；节奏由玩家行为自然形成。"),
        ("交付目标", "以一个公开链接完成访问，适配手机横屏，同时保留桌面浏览器演示能力。"),
        ("当前边界", "主体逻辑在客户端完成；暂未配置账号、数据库、排行榜与云端用户档案。"),
    ])

    add_heading(doc, "文档结构", 2)
    add_standard_table(doc, ["章节", "内容"], [
        ("1–2", "用户需求分析与设计概念"),
        ("3–4", "信息架构与完整交互流程"),
        ("5–6", "核心交互与界面视觉设计"),
        ("7", "行为数据和81种心签算法"),
        ("8–9", "技术实现路径、音效与性能策略"),
        ("10–12", "测试、发布、风险和后续迭代"),
    ], [1600, 7760], body_size=9.5)
    page_break(doc)


def add_user_requirements(doc: Document) -> None:
    add_heading(doc, "1. 用户需求分析", 1)
    add_heading(doc, "1.1 目标用户", 2)
    add_standard_table(doc, ["用户类型", "主要特征", "核心诉求"], [
        ("普通体验者", "对牧牛图了解有限，使用手机，注意力时间较短", "快速进入、操作直观、无需预备知识、能获得个人结果"),
        ("传统文化兴趣者", "关注石刻、禅宗图像和文化寓意", "内容有出处意识、视觉风格统一、互动不损害文化气质"),
        ("展览/课堂观众", "在展示现场或答辩中体验，时间有限", "流程稳定、无需登录、可快速复位、结果便于展示和讨论"),
        ("教师与评审", "关注设计逻辑、研究依据和技术完整性", "能说明需求、方法、算法、数据来源和实现边界"),
    ], [1900, 3200, 4260], body_size=8.8)

    add_heading(doc, "1.2 用户问题与设计响应", 2)
    add_standard_table(doc, ["用户问题", "需求转译", "设计响应"], [
        ("传统图像距离当代用户较远", "先理解再体验", "用石刻介绍、图像轮播和拓印建立真实文化入口"),
        ("纯文字讲解容易失去注意力", "将抽象禅理变成动作", "追、牵、稳、渡、收集、观照分别对应不同心性状态"),
        ("手机操作空间有限", "操作必须少而清楚", "采用点击、长按、拖动、双击等熟悉手势，并固定横屏舞台"),
        ("失败会打断沉浸", "过程比胜负重要", "游戏不存在永久失败；掉落后从头继续，累计时间作为记录"),
        ("结果若随机会削弱可信度", "个人结果需要可解释", "使用五项真实数据和确定性规则生成81种心签"),
        ("声音过多会破坏安静气质", "音效服务关键动作", "分阶段背景音乐，普通点击静音，仅保留必要动作反馈"),
    ], [2500, 2400, 4460], body_size=8.6)

    add_heading(doc, "1.3 功能需求", 2)
    add_text(doc, "系统需支持：完整文化介绍、三选一问心、拓印小游戏、四阶段连续流程、三类核心小游戏、实时数据采集、81种心签生成、手机横屏适配、背景音乐与动作音效、公开网址访问。")
    add_heading(doc, "1.4 非功能需求", 2)
    add_text(doc, "可用性方面，用户应在无额外说明的情况下理解主要手势；性能方面，资源需在移动网络下可加载且页面切换稳定；一致性方面，人物、牛、纸张、色彩与字体保持统一；可维护性方面，Web叙事层、Godot小游戏和算法模块相互独立；隐私方面，当前只在浏览器会话保存行为数据，不上传个人信息。")


def add_design_concept(doc: Document) -> None:
    add_heading(doc, "2. 设计概念", 1)
    add_heading(doc, "2.1 核心概念：牛即心，行即观", 2)
    add_callout(doc, "项目不是把传统图像做成装饰性小游戏，而是把牧牛图中的关系转译为交互关系：牛的奔逃对应心念散乱，绳索对应约束与觉照，独木桥对应持续调衡，莲花对应途中被看见的细微之物，最终心签则是对整段行为的文学化回望。")
    add_heading(doc, "2.2 四阶段体验结构", 2)
    add_standard_table(doc, ["阶段", "主题", "核心体验", "心理叙事"], [
        ("第一阶段", "寻牛 · 初识其心", "未牧六次点击、寻牛、受制长按", "先看见心的奔逸，再尝试以持续注意使其安定"),
        ("第二阶段", "驯牛 · 与心相处", "回首张力、独木桥、无碍对话", "从外力压制转向适度控制与相处"),
        ("第三阶段", "相忘 · 自在同行", "任运短动画、莲花跳跃、独照", "控制逐渐退场，人与牛进入自然同行"),
        ("第四阶段", "归真 · 回归本心", "心迹成图、心签呈现", "把全过程数据重新组织为个人化文字结果"),
    ], [1500, 1900, 2900, 3060], body_size=8.5)
    add_heading(doc, "2.3 设计原则", 2)
    add_key_value_table(doc, [
        ("文化先行", "先通过实物石刻、图像与拓印建立来源，再进入象征性交互。"),
        ("动作有义", "每个手势都对应一种精神关系，避免只为增加玩法而添加操作。"),
        ("无失败结局", "错误不终止体验，而被记录为用时、重试或注意力变化。"),
        ("渐进减法", "前期需要追逐和牵制，后期操作逐渐减少，直至观照与心签。"),
        ("结果可解释", "签文由明确数据产生，玩家看不到算法细节，但项目方可以完整说明。"),
    ])


def add_information_architecture(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "3. 信息架构", 1)
    add_heading(doc, "3.1 双层内容架构", 2)
    add_standard_table(doc, ["层级", "模块", "主要内容", "输出"], [
        ("文化认知层", "封面—问心—石刻介绍—十二境总览—拓印", "建立项目命题、选择初始牛、理解石刻来源与十二境含义", "初始路径 cowChoice"),
        ("互动体验层", "四阶段连续旅程", "将未牧、受制、回首、驯服、无碍、任运、相忘、独照等转化为互动", "抓牛、过桥、莲花等行为数据"),
        ("结果解释层", "心迹成图—心签", "将五项数据映射为路径、主导、节奏和层次", "81种正式心签之一"),
    ], [1600, 2500, 3560, 1700], body_size=8.5)

    add_heading(doc, "3.2 页面与状态结构", 2)
    add_text(doc, "Web主页面维护一个experience状态，决定当前展示介绍帧、阶段过渡页、CSS互动页或Godot iframe。介绍部分使用16张固定画面和热点区域；互动部分使用独立状态控制；小游戏通过统一URL参数进入对应Godot场景。")
    add_standard_table(doc, ["信息对象", "典型状态", "职责"], [
        ("介绍帧", "frames + screenIndex", "封面、问心、石刻与十二境介绍；负责文化叙事与导航"),
        ("CSS互动", "untrained / restrained / turning-back / free-roaming / unhindered / solitary", "承载轻量、节奏可控的状态动画与文字反馈"),
        ("Godot小游戏", "rubbing / seek / bridge / lotus", "承载碰撞、计时、收集、拖拽和角色运动"),
        ("阶段卡", "phase-one 至 phase-four", "以米色纸张、居中文字和莲纹完成章节过渡"),
        ("结果页", "heart-sign-forming", "读取数据、组合签文、以古籍竖排视觉展示"),
    ], [1900, 3000, 4460], body_size=8.7)
    add_figure(doc, WEB / "public/assets/heart-sign-rubbing-v2-landscape.png", "结果层视觉：横屏古籍式心签布局", 6.0)


def add_interaction_flow(doc: Document) -> None:
    add_heading(doc, "4. 完整交互流程", 1)
    add_callout(doc, "封面 → 命题介绍 → 问心三选一 → 石刻与十二境介绍 → 点击整页进入拓印 → 第一阶段 → 第二阶段 → 第三阶段 → 第四阶段 → 心签结果。各境互动完成后自动继续，不设置“循迹前行”等重复确认按钮。")
    add_standard_table(doc, ["序号", "节点", "用户动作", "系统反馈/去向"], [
        ("01", "封面", "点击“向心·开始旅程”", "进入项目命题介绍"),
        ("02", "问心", "在未牧、受制、任运三头牛中选择", "红色圈选反馈并记录cowChoice"),
        ("03", "文化介绍", "查看石刻、十二境含义与心月图", "通过画面热点顺序前行"),
        ("04", "拓印", "刷尘、覆纸、拍拓", "完成后进入第一阶段过渡页"),
        ("05", "未牧", "连续点击六次", "六帧逐步变化，牛挣脱向左出画"),
        ("06", "寻牛", "点击草地改变牧人方向，追上逃跑的牛", "60秒记录人物碰到牛的次数"),
        ("07", "受制", "持续长按2.8秒；中途松手需重来", "绳索绷紧、牛逐渐安定；完成后进入第二阶段"),
        ("08", "回首", "调节绳索张力并保持在40–60区间3秒", "太松/太紧都会回退，适中时牛主动回首"),
        ("09", "独木桥", "控制人物走到对岸", "掉落后从起点继续，累计总用时"),
        ("10", "无碍", "轻触画面两次查看问答", "第二次回答后自动进入第三阶段"),
        ("11", "任运", "观看快速帧动画", "人牛自在同行后自动进入莲花游戏"),
        ("12", "莲花跳跃", "轻触跳跃、双击高跳", "记录收集莲花数量和完成时间"),
        ("13", "独照", "观看笛音音符与提问，轻触继续", "进入第四阶段"),
        ("14", "心签", "等待边框与文字显现", "根据五项数据生成心签"),
    ], [620, 1350, 3370, 4020], body_size=8.1)


def add_core_interactions(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "5. 核心交互设计", 1)
    add_heading(doc, "5.1 拓印：把文化来源变成身体动作", 2)
    add_text(doc, "拓印不是装饰性开场，而是从“观看石刻”进入“亲手显影”的转换。流程按清尘、覆纸、拍拓三个步骤推进。刷尘与覆纸保留必要音效，连续拍拓声被取消，避免多次点击造成声音叠加。")
    add_heading(doc, "5.2 未牧与寻牛：从挣脱到追寻", 2)
    add_text(doc, "未牧使用七张状态图（初始加六次点击）形成逐帧CSS切换，第六次后牛向左跑出画面并自动进入寻牛。寻牛不设置输赢：60秒结束后保留抓到次数，玩家点击草地改变牧人方向，人物与牛发生碰撞即计数。")
    add_figure(doc, WEB / "public/assets/untrained-frame-0.png", "未牧互动初始画面：牛与牧人形成强烈对抗", 5.9)

    add_heading(doc, "5.3 受制与回首：从持续用力到适度张力", 2)
    add_text(doc, "受制要求长按约2.8秒。按住时牛的挣扎逐渐减弱；中途松手则进入“松绳，则野性复起”的回退状态。回首进一步把单纯长按改为张力调节：用户需要将张力保持在40–60的适中区间累计3秒；太松或太紧都会让进度回退。")
    add_figure(doc, WEB / "public/assets/turning-back-frame-return.png", "回首完成状态：牛第一次主动回望牧人", 5.9)

    add_heading(doc, "5.4 独木桥：失误不失败，时间成为记录", 2)
    add_text(doc, "玩家必须抵达对岸。掉落后人物回到起点，但计时继续，因此游戏结果不是“过或不过”，而是完成过程的总用时。实测无掉落约16秒，算法据此设置16秒、16–25秒、25秒以上三个区间。")
    add_heading(doc, "5.5 莲花跳跃：移动端节奏与观察", 2)
    add_text(doc, "角色自动向右推进，用户以轻触跳跃、双击高跳完成路线。人物比例、落脚线、平台碰撞和移动端视野均按横屏重构。游戏记录收集到的莲花数量和完成时间，分别用于判断观照程度与行动节奏。")


def add_visual_design(doc: Document) -> None:
    add_heading(doc, "6. 界面视觉设计", 1)
    add_heading(doc, "6.1 视觉方向", 2)
    add_callout(doc, "整体采用“纸本拓印 + 南宋牧牛图意象 + 当代横屏界面”的混合语言。画面不追求写实动画，而强调石刻、木版、水墨和旧纸纤维形成的手工痕迹。")
    add_standard_table(doc, ["视觉要素", "设计标准", "体验作用"], [
        ("色彩", "米白纸色为底，黛青为主文字与轮廓，朱砂红作选择和印章，少量橙金强调莲花与日月", "保持安静、古朴，同时让关键操作可识别"),
        ("字体", "中文界面统一使用汇文明朝体，正文保证可读字号", "形成古籍气质并维持全流程一致"),
        ("人物与牛", "固定牧人衣着、斗笠、体态和水牛角形；关键动作采用审核后的逐帧素材", "减少AI生成导致的形象漂移"),
        ("构图", "手机横屏16:9，主体位于中下区域，顶部保留标题或提示空间", "适应拇指操作并避免文字遮挡人物"),
        ("装饰", "莲花纹、云纹、印章与细线只作边界和章节提示", "建立文化识别，不挤占核心内容"),
        ("动效", "CSS位移、透明度、轻微抖动和帧切换为主", "保持拓印画面质感，避免现代化弹性动效破坏氛围"),
    ], [1700, 4400, 3260], body_size=8.4)
    add_heading(doc, "6.2 响应式与可用性", 2)
    add_text(doc, "基准设计画布为1066×600。页面根据浏览器可用宽高等比缩放并居中，手机端优先横屏。交互热点使用百分比定位，确保不同尺寸下仍与画面中的对象对齐；小游戏HUD采用放大字号和足够的边缘间距。")
    add_figure(doc, WEB / "public/assets/unhindered-dialogue.png", "无碍互动：对话位置避开人物与水牛主体", 6.0)


def add_data_algorithm(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "7. 数据与81种心签算法", 1)
    add_heading(doc, "7.1 五项真实输入", 2)
    add_standard_table(doc, ["数据", "来源", "用途"], [
        ("问心选择 cowChoice", "未牧/受制/任运三选一", "决定起心路径及并列分数优先级"),
        ("抓牛次数 caught", "寻牛游戏人物碰到牛的次数", "计算行动分并生成抓牛心解句"),
        ("过桥时间 bridgeSeconds", "真正到岸时的累计秒数", "计算平衡分并生成过桥心解句"),
        ("莲花数量 flowers", "莲花跳跃实际收集量", "计算观照分并生成莲花心解句"),
        ("莲花时间 lotusSeconds", "莲花跳跃抵达终点的总秒数", "判断疾、和、徐三种节奏"),
    ], [2400, 3600, 3360], body_size=8.7)

    add_heading(doc, "7.2 分数与分档", 2)
    add_standard_table(doc, ["维度", "算法", "当前阈值"], [
        ("行动分A", "clamp(抓牛次数 ÷ 18, 0, 1)", "0–9、10–17、18次及以上三档"),
        ("平衡分B", "按过桥总时间离散赋分", "≤16秒=1；17–25秒=0.65；≥26秒=0.30"),
        ("观照分C", "clamp(莲花数量 ÷ 48, 0, 1)", "0–19、20–34、35–48朵三档"),
        ("节奏", "按莲花完成时间分档", "≤65秒为疾；66–95秒为和；≥96秒为徐"),
        ("层次", "O=(A+B+C)÷3", "O<0.42初见；<0.72渐明；其余通透"),
    ], [1800, 3600, 3960], body_size=8.7)

    add_heading(doc, "7.3 81种组合逻辑", 2)
    add_callout(doc, "3种起心路径 × 3种主导行为 × 3种行动节奏 × 3种完成层次 = 81种正式心签。系统不使用随机数；相同输入必定得到相同结果。")
    add_text(doc, "主导行为取行动分、平衡分、观照分中最高者。若分数相同，则按问心选择设定稳定优先级：未牧优先行动，受制优先平衡，任运优先观照。每张正式签包含签名、四句签诗和七句心解。当前共有27种不同签名，层次通过第四句签诗和第七句心解继续区分，从而形成81个完整内容组合。")
    add_heading(doc, "7.4 解释边界", 2)
    add_text(doc, "心签描述的是本次游戏行为，不是心理量表，也不构成人格或医学诊断。阈值来自当前实测和设计判断，后续应通过更多玩家数据观察分布，再决定是否改为分位数校准。缺少任一数据时，系统显示已确认的备用“归径签”，避免生成不完整结果。", color=RED)


def add_technical_path(doc: Document) -> None:
    add_heading(doc, "8. 技术实现路径", 1)
    add_heading(doc, "8.1 总体架构", 2)
    add_standard_table(doc, ["层", "技术", "职责"], [
        ("Web叙事层", "React 19、Next 16兼容结构、Vinext、Vite 8、TypeScript", "页面状态、介绍帧、CSS互动、阶段过渡、音频控制和结果展示"),
        ("游戏运行层", "Godot 4.7.1 Web导出", "拓印、寻牛、独木桥、莲花跳跃的碰撞、计时和收集逻辑"),
        ("通信层", "iframe + window.postMessage", "Godot把完成事件、分数和音效提示传回Web主页面"),
        ("数据层", "React state + sessionStorage", "在当前浏览器会话保存五项数据并恢复状态"),
        ("规则层", "TypeScript确定性函数", "标准化数据、选择主导维度、生成81种心签"),
        ("资源层", "PNG/JPG/OTF/MP3", "逐帧人物与水牛、纸张纹理、字体、背景音乐和动作音效"),
    ], [1700, 3000, 4660], body_size=8.4)

    add_heading(doc, "8.2 Web与Godot通信", 2)
    add_text(doc, "主页面为每个Godot场景保留独立iframe引用，并校验消息来源，避免其他窗口伪造结果。Godot通过JavaScriptBridge向父页面发送结构化消息；Web接收后更新journeyData并推进experience状态。")
    add_standard_table(doc, ["消息", "数据", "作用"], [
        ("godot:stone-rubbing-complete", "无", "拓印完成，进入第一阶段"),
        ("godot:seek-score", "caught", "实时/结束时记录抓牛次数"),
        ("godot:bridge-complete", "seconds", "记录到岸累计时间并进入无碍"),
        ("godot:lotus-complete", "flowers, seconds", "记录莲花数量与完成时间并进入独照"),
        ("godot:audio", "cue", "触发经过白名单校验的动作音效"),
    ], [3000, 2100, 4260], body_size=8.7)

    add_heading(doc, "8.3 CSS互动与素材帧", 2)
    add_text(doc, "未牧、受制、回首、任运等段落不需要完整物理引擎，使用CSS动画和审核后的关键帧图片实现。这样能减少包体、控制节奏，并保持人物和水牛形象一致。涉及实时碰撞、角色运动、拖拽覆盖和计分的场景则交由Godot完成。")
    add_heading(doc, "8.4 数据持久化与隐私", 2)
    add_text(doc, "五项数据写入sessionStorage，仅在当前浏览器会话中使用。当前没有服务器数据库、用户账号或跨设备同步，因此项目严格来说是“前端应用 + 客户端游戏运行时”，而不是包含业务服务器的传统前后端系统。若未来增加排行榜或长期档案，需要新增后端API、数据库、用户同意和删除机制。")


def add_audio_performance(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "9. 声音、性能与移动端策略", 1)
    add_heading(doc, "9.1 声音设计", 2)
    add_standard_table(doc, ["类型", "当前策略"], [
        ("背景音乐", "按四个阶段切换，不使用一首音乐贯穿到底；首次用户触摸后解锁播放"),
        ("介绍页", "不播放翻页声和普通点击声，保持阅读安静"),
        ("拓印", "仅保留刷尘、覆纸等关键材质反馈，取消连续拍拓声叠加"),
        ("六次点击", "逐帧变化保持静音，避免按钮声破坏张力"),
        ("绳索", "拉紧音效限制约1.1秒，避免原始长音频覆盖后续画面"),
        ("小游戏", "抓牛、掉落、跳跃、收花等关键事件播放单次反馈"),
    ], [2000, 7360], body_size=9)
    add_heading(doc, "9.2 性能策略", 2)
    add_text(doc, "主要视觉资源在进入对应阶段时加载；Web层用图片和CSS承担轻量动画；Godot只在小游戏节点通过iframe运行。生产构建将Godot资源导出为Web包，静态文件与页面共同部署。后续上线前应继续压缩大型PNG与Godot PCK，检查移动网络首屏时间，并对低端Android设备进行内存测试。")
    add_heading(doc, "9.3 手机横屏策略", 2)
    add_text(doc, "用户通过手机浏览器打开公开网址，页面提示或默认采用横屏体验。所有重要文字和按钮避开屏幕边缘与系统手势区域；点击目标保持足够大；HUD字号按手机观看距离调整；角色和平台使用明确落脚线，避免视觉上悬空。")


def add_testing_delivery(doc: Document) -> None:
    add_heading(doc, "10. 测试与验收", 1)
    add_standard_table(doc, ["测试层级", "检查内容", "当前状态"], [
        ("单元测试", "流程映射、交互状态、算法边界、数据管线、音频路由", "自动化测试77项通过"),
        ("构建测试", "TypeScript编译、生产构建、静态资源路径", "生产构建通过"),
        ("Godot测试", "脚本解析、Web导出、场景消息", "Godot Web包已重新导出"),
        ("视觉验收", "人物/牛一致性、文字遮挡、落脚线、横屏缩放", "按阶段逐步由项目方人工确认"),
        ("流程验收", "从封面到心签不出现断链，小游戏完成可继续", "完整流程可在本地预览"),
        ("设备测试", "iOS Safari、Android Chrome、不同横屏尺寸", "发布前仍需真机矩阵复测"),
    ], [1800, 4600, 2960], body_size=8.7)
    add_heading(doc, "10.1 核心验收标准", 2)
    add_text(doc, "任何小游戏都不能形成不可恢复的死局；数据必须只记录真实行为；过桥掉落后计时继续；莲花同时记录数量与完成时间；心签同输入同输出；普通介绍页不产生密集点击声；字体、人物和水牛形象贯穿全流程一致。")

    add_heading(doc, "11. 发布与公开链接路径", 1)
    add_text(doc, "最终目标是将生产构建部署到支持静态站点或Node兼容构建的平台，获得HTTPS公开网址。用户无需安装应用，通过链接即可进入。部署流程为：完成生产构建 → 上传Web输出与Godot静态包 → 配置公开域名/平台域名 → 手机真机验收 → 将链接或二维码交付。")
    add_key_value_table(doc, [
        ("本地预览", "http://localhost:3002/（仅当前电脑可访问，不是公开网址）"),
        ("公开部署", "可选择Vercel、Cloudflare Pages或其他静态托管平台；需在最终发布阶段执行"),
        ("上线检查", "HTTPS、资源路径、缓存更新、横屏提示、移动端音频解锁、Godot加载进度"),
        ("交付形式", "一个公开链接，可另生成二维码用于展览和课堂展示"),
    ])


def add_risks_future(doc: Document) -> None:
    add_heading(doc, "12. 风险、限制与后续迭代", 1)
    add_standard_table(doc, ["问题", "当前限制", "建议"], [
        ("素材体积", "高清逐帧图片和Godot PCK较大", "WebP/AVIF压缩、按阶段预加载、拆分Godot场景包"),
        ("算法阈值", "部分阈值来自当前少量实测", "开展20–50人可用性测试，按真实分布校准"),
        ("设备差异", "不同手机浏览器的音频、触摸和内存表现不同", "建立iOS/Android真机测试矩阵"),
        ("结果解释", "文学化标签可能被误读为人格诊断", "结果页明确标注“本次行为映照”，答辩中说明边界"),
        ("数据保存", "sessionStorage关闭会话后消失", "如需长期档案，再增加用户授权、后端与隐私规则"),
        ("无障碍", "当前以视觉与触摸为主", "增加字幕、静音开关、键盘操作和高对比模式"),
    ], [1900, 3500, 3960], body_size=8.5)
    add_heading(doc, "结论", 1)
    add_callout(doc, "《牧心十二境》的价值不在于把传统文化包装成若干小游戏，而在于建立了一条完整的转译链：文化图像提供意义，交互动作让意义被身体感知，行为数据记录过程，心签再把过程返回给玩家。该结构使项目同时具备文化叙事、互动体验、可解释算法和公开发布的完整性。")

    add_heading(doc, "附录｜当前关键实现文件", 1)
    add_standard_table(doc, ["文件", "作用"], [
        ("app/page.tsx", "完整流程状态、数据收集、Godot消息、音频与页面呈现"),
        ("app/frame-flow.ts", "16张介绍帧、热点和导航关系"),
        ("app/player-sign.ts", "五项数据、81种心签算法和文字模块"),
        ("app/journey-audio.ts", "分阶段背景音乐与动作音效映射"),
        ("scripts/*.gd", "拓印、寻牛、独木桥和莲花游戏逻辑"),
        ("tests/*.test.mjs", "流程、算法、布局、互动和音频回归测试"),
    ], [3600, 5760], body_size=8.8)


def build() -> Path:
    doc = Document()
    configure(doc)
    build_cover(doc)
    add_executive_summary(doc)
    add_user_requirements(doc)
    add_design_concept(doc)
    add_information_architecture(doc)
    add_interaction_flow(doc)
    add_core_interactions(doc)
    add_visual_design(doc)
    add_data_algorithm(doc)
    add_technical_path(doc)
    add_audio_performance(doc)
    add_testing_delivery(doc)
    add_risks_future(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
